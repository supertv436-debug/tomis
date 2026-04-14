from flask import Flask, render_template, request, jsonify, send_file
import sqlite3
import os
import fitz
import docx
from docx.shared import Pt
import sys
import random
import re
import json
from pathlib import Path
from groq import Groq
from werkzeug.utils import secure_filename

app = Flask(__name__, template_folder='.')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['TEMPLATES_FOLDER'] = 'TEMPLATES'
app.config['EXPORTS_FOLDER'] = 'EXPORTS'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Создаем необходимые папки
for folder in [app.config['UPLOAD_FOLDER'], app.config['TEMPLATES_FOLDER'], app.config['EXPORTS_FOLDER']]:
    os.makedirs(folder, exist_ok=True)

# Инициализация БД
def init_db():
    conn = sqlite3.connect('tomis_memory.db')
    cur = conn.cursor()
    cur.execute("CREATE TABLE IF NOT EXISTS config (key TEXT PRIMARY KEY, val TEXT)")
    conn.commit()
    conn.close()

def get_api_key():
    conn = sqlite3.connect('tomis_memory.db')
    cur = conn.cursor()
    cur.execute("SELECT val FROM config WHERE key='api_key'")
    res = cur.fetchone()
    conn.close()
    return res[0] if res else None

def set_api_key(key):
    conn = sqlite3.connect('tomis_memory.db')
    cur = conn.cursor()
    cur.execute("INSERT OR REPLACE INTO config (key, val) VALUES ('api_key', ?)", (key,))
    conn.commit()
    conn.close()

def extract_text(file_path):
    try:
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            doc = fitz.open(file_path)
            return "\n".join(page.get_text() for page in doc)
        elif ext in ['.docx', '.doc']:
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        return ""
    except Exception as e:
        return f"ОШИБКА ЧТЕНИЯ: {e}"

def ask_ai(client, context, question, temperature=0.0):
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": "Ты строгий помощник кафедры. Отвечай точно по тексту силлабуса."},
                      {"role": "user", "content": f"КОНТЕКСТ:\n{context[:14000]}\n\nЗАДАНИЕ:\n{question}"}],
            temperature=temperature
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"ОШИБКА API: {str(e)}"

def fill_template_smart(template_path, output_path, data_dict):
    doc = docx.Document(template_path)
    clean_dict = {k.replace('*', '').strip().lower(): str(v) for k, v in data_dict.items()}

    def process_p(p):
        text = p.text
        orig = text
        for k, v in clean_dict.items():
            if k in text.lower():
                text = re.sub(re.escape(k), str(v), text, flags=re.IGNORECASE)
        if text != orig:
            text = re.sub(r'\s+\.', '.', text)
            text = re.sub(r' {2,}', ' ', text)
            p.text = text
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

    for p in doc.paragraphs:
        process_p(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_p(p)
    doc.save(output_path)

@app.route('/')
def index():
    with open('index.html', 'r', encoding='utf-8') as f:
        return f.read()

@app.route('/api/set_key', methods=['POST'])
def set_key():
    data = request.json
    api_key = data.get('api_key')
    if api_key and api_key.startswith('gsk_'):
        set_api_key(api_key)
        return jsonify({'success': True, 'message': 'API ключ сохранен'})
    return jsonify({'success': False, 'message': 'Неверный формат ключа'})

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Файл не выбран'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Файл не выбран'})
    
    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        text = extract_text(filepath)
        if "ОШИБКА" in text:
            return jsonify({'success': False, 'message': text})
        
        # Сохраняем текст в сессию или временный файл
        with open(os.path.join(app.config['UPLOAD_FOLDER'], 'current_context.txt'), 'w', encoding='utf-8') as f:
            f.write(text)
        
        return jsonify({'success': True, 'message': 'Файл успешно загружен', 'filename': filename})

@app.route('/api/ask', methods=['POST'])
def ask_question():
    api_key = get_api_key()
    if not api_key:
        return jsonify({'success': False, 'message': 'Сначала установите API ключ'})
    
    try:
        with open(os.path.join(app.config['UPLOAD_FOLDER'], 'current_context.txt'), 'r', encoding='utf-8') as f:
            context = f.read()
    except:
        return jsonify({'success': False, 'message': 'Сначала загрузите файл'})
    
    data = request.json
    question = data.get('question')
    
    client = Groq(api_key=api_key)
    answer = ask_ai(client, context, question)
    
    return jsonify({'success': True, 'answer': answer})

@app.route('/api/generate', methods=['POST'])
def generate_document():
    api_key = get_api_key()
    if not api_key:
        return jsonify({'success': False, 'message': 'Сначала установите API ключ'})

    data = request.json
    context = data.get('context', '')
    if not context:
        return jsonify({'success': False, 'message': 'Сначала загрузите файл'})

    doc_type = data.get('type')  # 'plan' or 'exam'
    
    # Получаем список шаблонов
    templates = [t for t in Path(app.config['TEMPLATES_FOLDER']).glob("*") if t.suffix.lower() in ['.docx', '.doc']]
    if not templates:
        return jsonify({'success': False, 'message': 'Папка TEMPLATES пуста'})
    
    template_path = None
    if doc_type == 'plan':
        template_path = next((t for t in templates if 'план' in t.name.lower() or 'calendar' in t.name.lower()), None)
    elif doc_type == 'exam':
        template_path = next((t for t in templates if 'билет' in t.name.lower() or 'ticket' in t.name.lower()), None)
    
    if not template_path:
        template_path = next((t for t in templates if t.suffix.lower() == '.docx'), templates[0])
    
    if doc_type == 'plan':
        task = """Изучи силлабус и извлеки данные для заполнения шаблона календарного плана.
Верни ТОЛЬКО валидный JSON (без markdown, без пояснений) со следующими ключами:
{
  "[ДИСЦИПЛИНА]": "название дисциплины из силлабуса",
  "[ГРУППЫ]": "номера групп или поток (если не указано — оставь пустым)",
  "[СЕМЕСТР]": "номер семестра (1 или 2)",
  "[ЧАСЫ_ЛЕК]": "количество часов лекций (число)",
  "[ЧАСЫ_ЛАБ]": "количество часов лабораторных/практических (число)",
  "[Л1]": "тема лекции 1",
  "[Л2]": "тема лекции 2",
  "[Л3]": "тема лекции 3",
  "[Л4]": "тема лекции 4",
  "[Л5]": "тема лекции 5",
  "[Л6]": "тема лекции 6",
  "[Л7]": "тема лекции 7",
  "[Л8]": "тема лекции 8",
  "[Л9]": "тема лекции 9",
  "[Л10]": "тема лекции 10",
  "[Л11]": "тема лекции 11",
  "[ПР1]": "тема лаб/практ работы 1",
  "[ПР2]": "тема лаб/практ работы 2",
  "[ПР3]": "тема лаб/практ работы 3",
  "[ПР4]": "тема лаб/практ работы 4",
  "[ПР5]": "тема лаб/практ работы 5",
  "[ПР6]": "тема лаб/практ работы 6",
  "[ПР7]": "тема лаб/практ работы 7",
  "[ПР8]": "тема лаб/практ работы 8"
}
Если тем меньше чем полей — заполни оставшиеся схожими темами из силлабуса."""
        fname = f"Календарный_план_{random.randint(10,99)}.docx"
        temp = 0.1
    else:
        task = """Изучи силлабус и сформируй данные для экзаменационных билетов.
Верни ТОЛЬКО валидный JSON (без markdown, без пояснений) со следующими ключами:
{
  "[КАФЕДРА]": "название кафедры из силлабуса",
  "[Дисциплина]": "название дисциплины",
  "[ШИФР]": "шифр специальности",
  "[СПЕЦ]": "название специальности",
  "[ЗАВ]": "ФИО заведующего кафедрой (если есть, иначе пустая строка)",
  "[ПРЕП]": "ФИО преподавателя (если есть, иначе пустая строка)",
  "[В1]": "вопрос 1",
  "[В2]": "вопрос 2",
  "[В3]": "вопрос 3",
  "[В4]": "вопрос 4",
  "[В5]": "вопрос 5",
  "[В6]": "вопрос 6",
  "[В7]": "вопрос 7",
  "[В8]": "вопрос 8",
  "[В9]": "вопрос 9",
  "[В10]": "вопрос 10",
  "[В12]": "вопрос 11",
  "[В13]": "вопрос 12",
  "[В14]": "вопрос 13",
  "[В15]": "вопрос 14",
  "[В16]": "вопрос 15",
  "[В17]": "вопрос 16",
  "[В18]": "вопрос 17",
  "[В19]": "вопрос 18",
  "[В20]": "вопрос 19",
  "[В21]": "вопрос 20",
  "[В22]": "вопрос 21",
  "[В23]": "вопрос 22",
  "[В24]": "вопрос 23",
  "[В25]": "вопрос 24",
  "[В26]": "вопрос 25",
  "[В27]": "вопрос 26",
  "[В28]": "вопрос 27",
  "[В29]": "вопрос 28",
  "[В30]": "вопрос 29",
  "[В31]": "вопрос 30",
  "[В32]": "вопрос 31",
  "[В33]": "вопрос 32",
  "[В34]": "вопрос 33"
}
Вопросы бери из раздела СРС, контрольных вопросов или экзаменационных вопросов силлабуса."""
        fname = f"Экзаменационные_билеты_{random.randint(10,99)}.docx"
        temp = 0.2
    
    client = Groq(api_key=api_key)
    ans = ask_ai(client, context, task, temp)
    
    # Извлекаем JSON из ответа (убираем markdown-блоки если есть)
    json_str = ans.strip()
    if '```' in json_str:
        import re as _re
        match = _re.search(r'```(?:json)?\s*([\s\S]*?)```', json_str)
        if match:
            json_str = match.group(1).strip()
    
    try:
        data_dict = json.loads(json_str)
    except:
        data_dict = {}

    # Для экзаменационных билетов — перемешиваем вопросы случайно без повторений
    if doc_type == 'exam' and data_dict:
        # Собираем все ключи вопросов в порядке [В1], [В2], ...
        question_keys = [k for k in data_dict if re.match(r'^\[В\d+\]$', k, re.IGNORECASE)]
        question_keys_sorted = sorted(question_keys, key=lambda k: int(re.search(r'\d+', k).group()))
        question_values = [data_dict[k] for k in question_keys_sorted]
        # Перемешиваем значения случайно
        random.shuffle(question_values)
        # Раскладываем обратно по тем же ключам
        for key, val in zip(question_keys_sorted, question_values):
            data_dict[key] = val

    output_path = os.path.join(app.config['EXPORTS_FOLDER'], fname)
    fill_template_smart(str(template_path), output_path, data_dict)
    
    return jsonify({'success': True, 'message': 'Документ создан', 'filename': fname})

GITHUB_RAW_BASE = "https://raw.githubusercontent.com/supertv436-debug/tomis/main/"

@app.route('/api/templates')
def get_templates():
    """Получить список всех шаблонов из templates.json с GitHub ссылками для скачивания"""
    try:
        with open('templates.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
        templates = data.get('templates', [])
        # Заменяем локальные пути на GitHub raw URL
        for t in templates:
            if 'filename' in t:
                t['filename'] = GITHUB_RAW_BASE + t['filename']
        return jsonify({'templates': templates})
    except Exception as e:
        return jsonify({'templates': [], 'error': str(e)})

@app.route('/TEMPLATES/<filename>')
def serve_template(filename):
    return send_file(os.path.join(app.config['TEMPLATES_FOLDER'], filename), as_attachment=True)

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(app.config['EXPORTS_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    init_db()
    app.run(debug=True, host='0.0.0.0', port=5000)
