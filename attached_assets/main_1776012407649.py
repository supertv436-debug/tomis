import sqlite3
import os
import fitz
import docx
from docx.shared import Pt
import sys
import time
import random
import re
import json
from pathlib import Path

# Для автоматической установки библиотек при первом запуске
try:
    from rich.prompt import Prompt
except ImportError:
    print("Установка необходимых библиотек...")
    os.system(f"{sys.executable} -m pip install rich groq tenacity python-docx pymupdf openpyxl")
    print("Библиотеки установлены. Перезапустите скрипт.")
    input("Нажмите Enter...")
    sys.exit()

import groq
from groq import Groq

os.system('')  # Включаем ANSI цвета в Windows

# ====================== ЦВЕТОВАЯ ПАЛИТРА КАРТУ ======================
C_CYAN = '\033[96m'
C_YELLOW = '\033[93m'
C_GREEN = '\033[92m'
C_RED = '\033[91m'
C_WHITE = '\033[97m'
C_GREY = '\033[90m'
C_RESET = '\033[0m'

def draw_logo():
    print(f"{C_CYAN}")
    print("   ============================================================================")
    print(f"          НАО «КАРАГАНДИНСКИЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ им. А. САГИНОВА»")
    print("   ============================================================================")
    print(f"      КАФЕДРА: {C_YELLOW}«Технологическое оборудование, машиностроение и стандартизация»{C_CYAN}")
    print("   ============================================================================")
    print(f"{C_YELLOW}")
    print("             ████████╗ ██████╗  ███╗   ███╗ ██╗ ███████╗")
    print("             ╚══██╔══╝██╔═══██╗ ████╗ ████║ ██║ ██╔════╝")
    print("                ██║   ██║   ██║ ██╔████╔██║ ██║ ███████╗")
    print("                ██║   ██║   ██║ ██║╚██╔╝██║ ██║ ╚════██║")
    print("                ██║   ╚██████╔╝ ██║ ╚═╝ ██║ ██║ ███████║")
    print("                ╚═╝    ╚═════╝  ╚═╝     ╚═╝ ╚═╝ ╚══════╝")
    print(f"{C_RESET}")

def loading_screen():
    os.system('cls' if os.name == 'nt' else 'clear')
    draw_logo()
    print(f"\n   {C_GREY}ИНИЦИАЛИЗАЦИЯ НЕЙРОСЕТИ КАФЕДРЫ (T.O.M.i.S. CORE v4.3)...{C_RESET}\n")
    modules = ["Ядро T.O.M.i.S.", "Подсистема Кафедры", "Smart Auto-Filler", "Поиск файлов по ПК"]
    for mod in modules:
        print(f"   {C_CYAN}[+]{C_RESET} Инициализация: {C_WHITE}{mod:<35}{C_RESET} {C_GREEN}[OK]{C_RESET}")
        time.sleep(0.25)
    print(f"\n   {C_GREEN}[✓] СИСТЕМА ГОТОВА. ДОБРО ПОЖАЛОВАТЬ, КИРИЛЛ (G.H.O.S.T.){C_RESET}\n")
    time.sleep(0.6)

def show_help():
    print(f"\n   {C_YELLOW}МАНУАЛ СИСТЕМЫ T.O.M.i.S. v4.3:{C_RESET}")
    help_data = [
        ("1. КЛЮЧ", "Установить Groq API ключ"),
        ("2. БАЗА", "Автоматический поиск + загрузка силлабуса"),
        ("3. ЭКСПОРТ", "Вырезка данных из Excel"),
        ("4. БЛАНКИ", "Открыть шаблоны из папки TEMPLATES"),
        ("5. ВОПРОС", "Диалог с ИИ по загруженному файлу"),
        ("6. ПЛАН", "Авто-генерация календарного плана"),
        ("7. ЭКЗАМЕН", "Генерация экзаменационных билетов"),
        ("8. МЕНЮ", "Очистка и обновление экрана"),
        ("9. ИНФО", "Эта справка"),
        ("0. ВЫХОД", "Завершить работу")
    ]
    for cmd, desc in help_data:
        print(f"   {C_CYAN}{cmd:<12}{C_RESET} — {C_GREY}{desc}{C_RESET}")
    print()

def draw_interface(api_status, db_status):
    os.system('cls' if os.name == 'nt' else 'clear')
    draw_logo()
    print(f"   {C_WHITE}Кафедральный AI-ассистент v4.3{C_RESET}  |  {C_YELLOW}Силин Кирилл{C_RESET}")
    print(f"   {C_GREY}────────────────────────────────────────────────────────────────────────────────────────{C_RESET}")
    print(f"   {C_CYAN}СИСТЕМНЫЙ СТАТУС:{C_RESET}   API CORE: {api_status}   |   DATA MEMORY: {db_status}")
    print(f"   {C_GREY}────────────────────────────────────────────────────────────────────────────────────────{C_RESET}")
    print(f"   {C_CYAN}[1]{C_RESET} {C_WHITE}КЛЮЧ{C_RESET}     {C_GREY}- Установить API-ключ Groq{C_RESET}")
    print(f"   {C_CYAN}[2]{C_RESET} {C_WHITE}БАЗА{C_RESET}     {C_GREY}- Автопоиск силлабуса по всему ПК{C_RESET}")
    print(f"   {C_CYAN}[3]{C_RESET} {C_WHITE}ЭКСПОРТ{C_RESET}  {C_GREY}- Вырезать данные из Excel{C_RESET}")
    print(f"   {C_CYAN}[4]{C_RESET} {C_WHITE}БЛАНКИ{C_RESET}   {C_GREY}- Открыть шаблон Word{C_RESET}")
    print(f"   {C_CYAN}[5]{C_RESET} {C_WHITE}ВОПРОС{C_RESET}   {C_GREY}- Задать вопрос ИИ{C_RESET}")
    print(f"   {C_CYAN}[6]{C_RESET} {C_WHITE}ПЛАН{C_RESET}     {C_GREY}- Генерация календарного плана{C_RESET}")
    print(f"   {C_CYAN}[7]{C_RESET} {C_WHITE}ЭКЗАМЕН{C_RESET}  {C_GREY}- Генерация билетов из СРС{C_RESET}")
    print(f"   {C_CYAN}[8]{C_RESET} {C_WHITE}МЕНЮ{C_RESET}     {C_GREY}- Обновить экран{C_RESET}")
    print(f"   {C_CYAN}[9]{C_RESET} {C_WHITE}ИНФО{C_RESET}     {C_GREY}- Показать справку{C_RESET}")
    print(f"   {C_GREY}[0] ВЫХОД     - Закрыть систему{C_RESET}")
    print(f"   {C_GREY}━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━{C_RESET}\n")

def extract_text(file_path):
    try:
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            doc = fitz.open(file_path)
            return "\n".join(page.get_text() for page in doc)
        elif ext in ['.docx', '.doc']:
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        return ""
    except Exception as e:
        return f"ОШИБКА ЧТЕНИЯ: {e}"

def find_syllabus_files():
    """Автоматический поиск силлабусов по всему ПК"""
    print(f"   {C_CYAN}[+] Запуск поиска силлабусов по компьютеру...{C_RESET}")
    
    # 1. Сначала проверяем папку проекта
    base_p = Path(os.path.dirname(os.path.abspath(sys.argv[0])))
    candidates = list(base_p.rglob("*.pdf")) + list(base_p.rglob("*.docx"))
    
    # 2. Популярные папки пользователя
    user_dirs = [
        Path.home() / "Desktop",
        Path.home() / "Рабочий стол",
        Path.home() / "Documents",
        Path.home() / "Документы",
        Path.home() / "Downloads",
        Path.home() / "Загрузки",
        Path.home() / "OneDrive",
    ]
    
    for dir_path in user_dirs:
        if dir_path.exists():
            candidates.extend(list(dir_path.rglob("*.pdf")))
            candidates.extend(list(dir_path.rglob("*.docx")))
    
    # Убираем дубликаты и сортируем по имени
    unique_files = []
    seen = set()
    for f in candidates:
        if f.name.lower() not in seen and any(word in f.name.lower() for word in ["силлабус", "syllabus", "программа", "дисциплин", "рабочей"]):
            seen.add(f.name.lower())
            unique_files.append(f)
    
    # Если ничего релевантного не нашли — показываем все PDF и DOCX
    if not unique_files:
        unique_files = list(dict.fromkeys(candidates))  # убираем дубли
    
    return unique_files[:20]  # ограничиваем 20 файлами для удобства

def ask_ai(client, context, question, temperature=0.0):
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": "Ты строгий помощник кафедры. Отвечай точно по тексту силлабуса."},
                      {"role": "user", "content": f"КОНТЕКСТ:\n{context[:14000]}\n\nЗАДАНИЕ:\n{question}"}],
            temperature=temperature
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"ОШИБКА API: {str(e)}"

def fill_template_smart(template_path, output_path, data_dict):
    doc = docx.Document(template_path)
    clean_dict = {k.replace('*', '').strip().lower(): str(v) for k, v in data_dict.items()}
    
    def process_p(p):
        text = p.text
        orig = text
        for k, v in clean_dict.items():
            if k in text.lower():
                text = re.sub(re.escape(k), str(v), text, flags=re.IGNORECASE)
        if text != orig:
            text = re.sub(r'\s+\.', '.', text)
            text = re.sub(r' {2,}', ' ', text)
            p.text = text
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
    
    for p in doc.paragraphs: process_p(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: process_p(p)
    doc.save(output_path)

def main():
    loading_screen()
    
    base_p = Path(os.path.dirname(os.path.abspath(sys.argv[0])))
    db_path = base_p / "tomis_memory.db"
    t_dir = base_p / "TEMPLATES"
    e_dir = base_p / "EXPORTS"
    t_dir.mkdir(exist_ok=True)
    e_dir.mkdir(exist_ok=True)
    
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("CREATE TABLE IF NOT EXISTS config (key TEXT PRIMARY KEY, val TEXT)")
    conn.commit()
    
    cur.execute("SELECT val FROM config WHERE key='api_key'")
    res = cur.fetchone()
    api_key = res[0] if res else None
    client = Groq(api_key=api_key) if api_key else None
    
    doc_ctx = ""
    chat_history = []
    
    api_status = f"{C_GREEN}УСТАНОВЛЕН{C_RESET}" if api_key else f"{C_RED}ОТСУТСТВУЕТ{C_RESET}"
    db_status = f"{C_RED}ПУСТО{C_RESET}"
    
    draw_interface(api_status, db_status)
    
    while True:
        cmd = input(f"   {C_YELLOW}ВЫБОР КОМАНДЫ > {C_RESET}").strip()
        
        if cmd == '0':
            break
            
        elif cmd == '8':
            draw_interface(api_status, db_status)
            
        elif cmd == '9':
            show_help()
            
        elif cmd == '1':
            k = input(f"   {C_GREY}► Введите Groq API ключ (gsk_...): {C_RESET}").strip()
            if k.startswith("gsk_"):
                cur.execute("INSERT OR REPLACE INTO config (key, val) VALUES ('api_key', ?)", (k,))
                conn.commit()
                api_key = k
                client = Groq(api_key=api_key)
                api_status = f"{C_GREEN}УСТАНОВЛЕН{C_RESET}"
                print(f"   {C_GREEN}[✓] Ключ сохранён{C_RESET}\n")
            else:
                print(f"   {C_RED}[!] Ключ должен начинаться с gsk_{C_RESET}\n")
                
        elif cmd == '2':   # ← Главное изменение здесь
            files = find_syllabus_files()
            
            if not files:
                print(f"   {C_RED}[!] Файлы не найдены на компьютере.{C_RESET}")
                f_path = input(f"   {C_GREY}► Введите путь вручную: {C_RESET}").strip('"\'')
            else:
                print(f"\n   {C_YELLOW}Найдено файлов: {len(files)}{C_RESET}")
                for i, f in enumerate(files, 1):
                    print(f"     {C_CYAN}[{i}]{C_RESET} {f.name}  {C_GREY}({f.parent.name}){C_RESET}")
                
                choice = input(f"\n   {C_GREY}► Выберите номер файла или введите путь вручную: {C_RESET}").strip()
                
                if choice.isdigit() and 1 <= int(choice) <= len(files):
                    f_path = str(files[int(choice)-1])
                else:
                    f_path = choice.strip('"\'')
            
            if os.path.exists(f_path):
                print(f"   {C_CYAN}[+] Чтение силлабуса...{C_RESET}")
                doc_ctx = extract_text(f_path)
                if "ОШИБКА" not in doc_ctx:
                    chat_history = []
                    db_status = f"{C_GREEN}ГОТОВА{C_RESET}"
                    print(f"   {C_GREEN}[+] Файл успешно загружен!{C_RESET}\n")
                else:
                    print(f"   {C_RED}[!] {doc_ctx}{C_RESET}\n")
            else:
                print(f"   {C_RED}[!] Файл не найден по указанному пути.{C_RESET}\n")
                
        elif cmd in ['6', '7']:
            if not client or not doc_ctx:
                print(f"   {C_RED}[!] Сначала установите ключ [1] и загрузите базу [2]{C_RESET}\n")
                continue
                
            templates = list(t_dir.glob("*.docx"))
            if not templates:
                print(f"   {C_RED}[!] Положите .docx шаблон в папку TEMPLATES{C_RESET}\n")
                continue
                
            print(f"   {C_YELLOW}Шаблоны:{C_RESET}")
            for i, t in enumerate(templates, 1):
                print(f"     {C_CYAN}[{i}]{C_RESET} {t.name}")
            
            t_idx = input(f"   {C_GREY}► Номер шаблона: {C_RESET}").strip()
            if not t_idx.isdigit() or int(t_idx) < 1 or int(t_idx) > len(templates):
                continue
            sel_t = templates[int(t_idx)-1]
            
            if cmd == '6':
                task = "Изучи силлабус и верни данные для календарного плана в формате JSON."
                fname = f"План_ГОТОВЫЙ_{random.randint(10,99)}.docx"
                temp = 0.2
            else:
                task = "Найди все контрольные вопросы для СРС и верни их в JSON формате."
                fname = f"Билеты_ГОТОВЫЕ_{random.randint(10,99)}.docx"
                temp = 0.6
            
            print(f"   {C_CYAN}[+] Генерация документа...{C_RESET}")
            ans = ask_ai(client, doc_ctx, task, temp)
            
            try:
                data = json.loads(ans)
            except:
                data = {}
                for line in ans.split('\n'):
                    if ':::' in line:
                        k, v = line.split(':::', 1)
                        data[k.strip().upper()] = v.strip()
            
            output_path = base_p / fname
            fill_template_smart(str(sel_t), str(output_path), data)
            
            print(f"   {C_GREEN}[✓] Документ создан: {fname}{C_RESET}")
            try:
                os.startfile(str(output_path))
            except:
                print(f"   Сохранён по пути: {output_path}")
        
        elif cmd == '4':
            templates = list(t_dir.glob("*.docx"))
            if templates:
                for i, t in enumerate(templates, 1):
                    print(f"   {C_CYAN}[{i}]{C_RESET} {t.name}")
                idx = input(f"   {C_GREY}► Номер: {C_RESET}").strip()
                if idx.isdigit() and 1 <= int(idx) <= len(templates):
                    os.startfile(str(templates[int(idx)-1]))
            else:
                print(f"   {C_RED}[!] Папка TEMPLATES пуста{C_RESET}\n")
        
        else:
            print(f"   {C_RED}[!] Неизвестная команда. Введите 9 для справки.{C_RESET}\n")

    conn.close()
    print(f"\n   {C_CYAN}До свидания, Кирилл!{C_RESET}")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(f"\n\n   {C_RED}[!] Работа прервана.{C_RESET}")
    except Exception as e:
        print(f"\n   {C_RED}[КРИТИЧЕСКАЯ ОШИБКА] {e}{C_RESET}")
        input("Нажмите Enter...")